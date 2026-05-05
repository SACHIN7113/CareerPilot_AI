import asyncio
from io import BytesIO

import docx

from app.services.document_parser import chunk_text_async, extract_text_async


def _run(coro):
    return asyncio.run(coro)


def _build_docx_bytes() -> bytes:
    document = docx.Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Company Name"
    table.cell(0, 1).text = "Acme Corp"
    table.cell(1, 0).text = "Role"
    table.cell(1, 1).text = "Software Developer"
    document.add_paragraph("We are hiring a Software Developer for backend systems.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_chunk_text_keeps_question_and_answer_lines_together_when_possible() -> None:
    text = "\n".join(
        [
            "Q1: What is Python?",
            "Answer: Python is a high-level programming language.",
            "",
            "Q2: What is Pandas?",
            "Answer: Pandas is a data analysis library.",
        ]
    )

    chunks = _run(chunk_text_async(text, chunk_size=120, overlap=30))

    assert len(chunks) >= 2
    assert "Q1: What is Python?" in chunks[0]
    assert "Answer: Python is a high-level programming language." in chunks[0]


def test_chunk_text_splits_long_plain_paragraphs_without_returning_empty_chunks() -> None:
    text = " ".join(["Python helps with automation and analysis."] * 80)

    chunks = _run(chunk_text_async(text, chunk_size=180, overlap=40))

    assert chunks
    assert all(chunk.strip() for chunk in chunks)
    assert all(len(chunk) <= 220 for chunk in chunks)


def test_extract_text_includes_docx_tables() -> None:
    text = _run(extract_text_async("job_description.docx", _build_docx_bytes()))

    assert "Acme Corp" in text
    assert "Software Developer" in text
